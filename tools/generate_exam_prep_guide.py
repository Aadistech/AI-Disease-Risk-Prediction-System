from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = PROJECT_ROOT / "reports"
OUTPUT_DOCX = REPORTS_DIR / "AI_Multi_Disease_Exam_Prep_Guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Powered Multi-Disease Risk Prediction System")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Exam Preparation Guide for HOD and External Viva")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Prepared from the project files in this repository")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    set_cell_margins(table)

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], LIGHT_BLUE)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table, top=120, bottom=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(body)


def build_document():
    REPORTS_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "One-line project explanation",
        "This project uses machine learning to predict diabetes and heart disease risk from patient health parameters, then presents real-time predictions through a Streamlit dashboard.",
    )

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "The project is an AI-powered multi-disease risk prediction system. It currently supports two diseases: diabetes and heart disease. The training work is done in Jupyter notebooks, the trained models are saved as pickle files, and the dashboard loads those models for live prediction."
    )
    add_bullets(
        doc,
        [
            "Problem area: healthcare analytics and early risk screening.",
            "Input: patient medical measurements such as glucose, BMI, blood pressure, cholesterol, age, and heart-specific clinical indicators.",
            "Output: binary prediction, where 1 means higher risk and 0 means lower risk.",
            "Deployment: Streamlit web application in dashboard/app.py.",
            "Important limitation: this is not a replacement for a doctor; it is a decision-support and awareness tool.",
        ],
    )

    doc.add_heading("2. Project Structure", level=1)
    add_table(
        doc,
        ["File or Folder", "Purpose"],
        [
            ["README.md", "Project overview, setup instructions, technologies, workflow, and future scope."],
            ["requirements.txt", "Lists Python libraries needed to run the project."],
            ["dashboard/app.py", "Main Streamlit dashboard for analytics and prediction."],
            ["data/diabetes.csv", "Diabetes dataset with 768 rows, 8 input features, and Outcome as target."],
            ["data/heart.csv", "Heart disease dataset with 1025 rows, 13 input features, and target as output."],
            ["models/diabetes_model.pkl", "Saved diabetes prediction model loaded by the dashboard."],
            ["models/heart_model.pkl", "Saved heart disease prediction model loaded by the dashboard."],
            ["notebooks/eda.ipynb", "Diabetes EDA, model training, evaluation, and model saving."],
            ["notebooks/heart_disease_model.ipynb", "Heart disease model training, evaluation, and saving workflow."],
            ["reports/", "Contains existing project reports and this viva preparation guide."],
        ],
        widths=[2.0, 4.5],
    )

    doc.add_heading("3. Technologies Used", level=1)
    add_table(
        doc,
        ["Technology", "Use in Project"],
        [
            ["Python", "Main programming language."],
            ["Pandas", "Reads CSV files and creates DataFrames for model input."],
            ["Matplotlib and Seaborn", "Used in notebooks for charts, count plots, heatmaps, and EDA."],
            ["Scikit-learn", "Used for train-test split, model training, prediction, and evaluation."],
            ["Logistic Regression", "Main binary classification model described in the notebooks."],
            ["Random Forest", "Compared in the diabetes notebook."],
            ["Joblib", "Saves and loads trained machine learning models."],
            ["Streamlit", "Creates the interactive web dashboard."],
        ],
        widths=[1.8, 4.7],
    )

    doc.add_heading("4. Machine Learning Workflow", level=1)
    add_numbered(
        doc,
        [
            "Data collection: diabetes.csv and heart.csv are used as structured healthcare datasets.",
            "Data preprocessing: the notebooks inspect data information, missing values, and statistical summary.",
            "Exploratory data analysis: charts and correlation heatmaps are created to understand patterns.",
            "Feature selection: input columns are stored in X and target columns are stored in y.",
            "Train-test split: the data is divided into training and testing sets using train_test_split.",
            "Model training: Logistic Regression is trained, and diabetes also compares Random Forest.",
            "Model evaluation: accuracy, confusion matrix, and classification report are used.",
            "Model saving: joblib.dump saves the trained models into the models folder.",
            "Deployment: Streamlit loads the saved models and predicts risk from user input.",
        ],
    )

    doc.add_heading("5. Dashboard Code: Line-by-Line Speaking Guide", level=1)
    add_table(
        doc,
        ["Lines", "What to Say"],
        [
            ["1-7", "These lines import required libraries: Path for paths, joblib for loading models, pandas for data handling, Streamlit for UI, and sklearn utilities."],
            ["10-14", "These lines define the project root and build paths for diabetes data, diabetes model, and heart model."],
            ["16-24", "These lists store the exact feature names expected by each trained model. Feature order must match the training data."],
            ["27-39", "These functions load the dataset and saved models. Streamlit caching improves performance by avoiding repeated loading."],
            ["42-51", "This configures the Streamlit page title, page icon, layout, and main heading."],
            ["55-58", "The selectbox lets the user choose between Diabetes and Heart Disease prediction."],
            ["62-64", "If Diabetes is selected, the app loads diabetes data and the saved diabetes model."],
            ["66-73", "The sidebar shows basic dataset metrics such as total rows, features, diabetic cases, and non-diabetic cases."],
            ["75-96", "The main diabetes analytics area shows patient metrics, class distribution, and BMI vs glucose scatter chart."],
            ["100-114", "These number inputs collect diabetes patient values from the user."],
            ["116-133", "After clicking the button, the app creates a one-row DataFrame, predicts risk, shows probability if available, and displays a recommendation."],
            ["135-137", "The app displays the first 20 rows of the diabetes dataset for transparency."],
            ["140-144", "If Heart Disease is selected, the app loads the saved heart model and heart dataset."],
            ["146-157", "The heart disease sidebar shows total rows, features, disease cases, and no-disease cases."],
            ["159-191", "The heart analytics area shows metrics, target distribution, and age vs cholesterol scatter chart."],
            ["195-216", "These inputs collect 13 heart disease features from the user."],
            ["218-253", "The app creates a heart patient DataFrame, predicts risk, shows probability, and displays risk-specific guidance."],
            ["255-262", "The app displays a preview of the heart dataset."],
        ],
        widths=[1.0, 5.5],
    )

    doc.add_heading("6. Diabetes Dataset Explanation", level=1)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["Pregnancies", "Number of pregnancies."],
            ["Glucose", "Blood glucose level."],
            ["BloodPressure", "Blood pressure value."],
            ["SkinThickness", "Skin fold thickness."],
            ["Insulin", "Insulin level."],
            ["BMI", "Body mass index."],
            ["DiabetesPedigreeFunction", "Family history or genetic tendency toward diabetes."],
            ["Age", "Patient age."],
            ["Outcome", "Target value: 1 means diabetic, 0 means non-diabetic."],
        ],
        widths=[2.0, 4.5],
    )

    doc.add_heading("7. Heart Disease Dataset Explanation", level=1)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["age", "Patient age."],
            ["sex", "0 means female, 1 means male."],
            ["cp", "Chest pain type."],
            ["trestbps", "Resting blood pressure."],
            ["chol", "Cholesterol level."],
            ["fbs", "Fasting blood sugar greater than 120 mg/dl."],
            ["restecg", "Resting ECG result."],
            ["thalach", "Maximum heart rate achieved."],
            ["exang", "Exercise-induced angina."],
            ["oldpeak", "ST depression value."],
            ["slope", "Slope of the ST segment."],
            ["ca", "Number of major vessels."],
            ["thal", "Thalassemia category."],
            ["target", "Target value: 1 means higher heart disease risk, 0 means lower risk."],
        ],
        widths=[1.6, 4.9],
    )

    doc.add_heading("8. Notebook Explanation", level=1)
    doc.add_paragraph(
        "Both notebooks follow the same core machine learning lifecycle. The diabetes notebook also compares Logistic Regression with Random Forest. The notebook conclusion states that Logistic Regression achieved approximately 74.67 percent accuracy and was selected as the final diabetes model."
    )
    add_bullets(
        doc,
        [
            "eda.ipynb: diabetes EDA, visualization, Logistic Regression, Random Forest comparison, confusion matrix, classification report, and joblib model saving.",
            "heart_disease_model.ipynb: heart disease EDA, Logistic Regression training, prediction, accuracy, confusion matrix, classification report, model saving, and model loading test.",
        ],
    )

    doc.add_heading("9. Two-Minute Presentation Script", level=1)
    script = [
        "Good morning respected HOD, external examiner, and faculty members. My project is titled AI-Powered Multi-Disease Risk Prediction System. The aim of this project is to use machine learning to predict disease risk at an early stage using patient medical parameters.",
        "Currently, the system supports diabetes and heart disease prediction. I collected structured healthcare datasets, performed exploratory data analysis, selected input features, trained machine learning models, evaluated them using accuracy, confusion matrix, and classification report, and then deployed the saved models through a Streamlit web application.",
        "For diabetes, the model uses features like glucose, BMI, blood pressure, insulin, age, and family history. For heart disease, it uses features like age, cholesterol, resting blood pressure, chest pain type, ECG, and maximum heart rate. The user enters these values in the dashboard, and the system predicts whether the patient has higher or lower disease risk.",
        "The main advantage of this project is that it demonstrates how machine learning can support early healthcare screening. It is not a replacement for doctors, but it can act as a decision-support tool for awareness and preliminary risk analysis.",
    ]
    for paragraph in script:
        doc.add_paragraph(paragraph)

    doc.add_heading("10. Likely Viva Questions and Answers", level=1)
    qa = [
        ("What is the objective of your project?", "To predict diabetes and heart disease risk using machine learning and provide a simple dashboard for real-time prediction."),
        ("Which algorithm did you use?", "Logistic Regression is the main final model described in the notebooks. Random Forest is also compared in the diabetes notebook."),
        ("Why Logistic Regression?", "Because the target output is binary, either risk or no risk. Logistic Regression is simple, interpretable, and suitable for binary classification."),
        ("What is Streamlit?", "Streamlit is a Python framework used to create interactive web apps for data science and machine learning projects."),
        ("What is joblib used for?", "It is used to save and load trained machine learning models as pickle files."),
        ("What is the target column?", "For diabetes, the target is Outcome. For heart disease, the target is target."),
        ("Is this a medical diagnosis system?", "No. It is a risk prediction and decision-support system. Final diagnosis must be done by medical professionals."),
        ("What are the limitations?", "The prediction depends on dataset quality, model accuracy, and entered values. It does not replace expert medical diagnosis."),
        ("What is future scope?", "Improve model accuracy, add more diseases, deploy online, add patient login and history, and use more advanced models."),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        q = p.add_run("Q. " + question + " ")
        q.bold = True
        doc.add_paragraph("A. " + answer, style=None)

    doc.add_heading("11. Important Correction to Remember", level=1)
    add_callout(
        doc,
        "README vs code",
        "The README mentions that the app can retrain a Random Forest model if the saved diabetes model is missing, but the current dashboard code directly loads the saved model files. If asked, explain that fallback retraining is documented but not implemented in the current app.py.",
    )

    doc.add_heading("12. Final Closing Line", level=1)
    doc.add_paragraph(
        "In conclusion, this project demonstrates the complete machine learning pipeline: data analysis, model training, model evaluation, model saving, and real-time deployment through a web dashboard."
    )

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_document())
